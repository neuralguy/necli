"""create_docx — конвертация HTML в DOCX через Pandoc + post-processing.

Возможности:
  • LaTeX-математика ($...$, $$...$$) → нативные формулы Word (OMML)
  • Inline-CSS (color, font-family, font-size, background-color) применяется
    через post-processing python-docx (pandoc HTML reader их не понимает)
  • Картинки: data URI или локальные пути
  • Дефолтный reference.docx со стилями: TNR 14, межстрочный 1.5,
    чёрные заголовки

Внешние зависимости: pandoc 3.x в PATH, python-docx.
"""

from __future__ import annotations

import re
import subprocess
import uuid
from logger import logger
from tools._paths import resolve_path, clean_path, get_working_dir
from tools.models import ToolCall, ToolResult
from tools.file_ops._docx_reference import get_default_reference_path
from tools.file_ops._html_preprocess import (
    wrap_table_cells as _wrap_table_cells,
    extract_code_blocks as _extract_code_blocks,
    restore_code_blocks as _restore_code_blocks,
    extract_stray_angles as _extract_stray_angles,
    restore_stray_angles as _restore_stray_angles,
)
from tools.file_ops._pandoc import find_pandoc as _find_pandoc, install_hint as _install_hint
from tools.file_ops._docx_sources import (
    save_source as _save_docx_source,
    load_template as _load_docx_template,
    iter_templates as _iter_docx_templates,
)


# Конвертация CSS-единиц шрифта в pt (python-docx Pt):
#   1px ≈ 0.75pt (96 dpi: 1pt = 1/72in, 1px = 1/96in → 72/96 = 0.75);
#   1em ≈ 14pt   (базовый размет документа = TNR 14, см. _docx_reference).
_PX_TO_PT = 0.75
_EM_TO_PT = 14

# Ширина текстовой колонки страницы A4 в twips: 21cm − поля 2×2cm = 17cm.
# 17cm × 567 twips/cm ≈ 9639; используем округлённое значение для tblGrid.
_PAGE_TEXT_WIDTH_TWIPS = 9360

_MARKER_RE = re.compile(r"^\[DOCX as HTML[^\]]*\]\s*\n?", flags=re.IGNORECASE)


def _strip_read_marker(content: str) -> str:
    return _MARKER_RE.sub("", content, count=1)


def _find_matching_template(out_path, content: str):
    """Ищет шаблон-оригинал для точного round-trip.

    Приоритет:
      1. Шаблон, ключённый по самому out_path (перезапись того же файла).
      2. Любой сохранённый шаблон, чей текст хорошо выравнивается с content
         (агент прочитал X.docx и пишет в X_копия.docx — другой путь).
    Возвращает Path к .template.docx или None.
    """
    tpl = _load_docx_template(out_path)
    if tpl is not None:
        return tpl

    try:
        from docx import Document
        from tools.file_ops._docx_whitespace import (
            tree_walk_paragraphs, parse_html_blocks, collapse_ws,
        )
        import difflib

        want_texts = [
            collapse_ws(b["text"]) for b in parse_html_blocks(content)
            if b["text"].strip()
        ]
        if not want_texts:
            return None

        best = None
        best_ratio = 0.0
        for cand in _iter_docx_templates():
            try:
                doc = Document(str(cand))
            except Exception:
                continue
            tpl_texts = [
                collapse_ws(p.text) for p in tree_walk_paragraphs(doc)
                if p.text.strip()
            ]
            if not tpl_texts:
                continue
            ratio = difflib.SequenceMatcher(
                None, want_texts, tpl_texts, autojunk=False
            ).ratio()
            if ratio > best_ratio:
                best_ratio, best = ratio, cand
        # Берём шаблон только при уверенном совпадении содержимого.
        if best is not None and best_ratio >= 0.85:
            return best
    except Exception:
        logger.debug("template match by content skipped", exc_info=True)
    return None


_MATH_BLOCK_RE = re.compile(r"(\$\$[\s\S]+?\$\$|(?<!\$)\$[^\$\n]+?\$(?!\$))")

# Текстовые/операторные аргументы, внутри которых скобки трогать нельзя:
# \text{(note)}, \mathrm{(p)}, \operatorname{...} и т.п.
_TEXT_ARG_RE = re.compile(
    r"\\(?:text|textrm|textit|textbf|textsf|texttt|mathrm|mathsf|mathtt|mathit|mathbf|operatorname|operatorname\*)\s*\{",
)


def _fix_math_parens(content: str) -> str:
    """Заменяет голые ( и ) внутри $...$ / $$...$$ на \\left( \\right).

    LibreOffice OMML renderer плохо рендерит обычные скобки вокруг
    оснований степеней — (1+y)^{...} ломается. \\left/\\right даёт OMML
    <m:d>, который понимает и Word, и LibreOffice.

    Защиты:
      - не трогаем уже завёрнутые \\left( / \\right);
      - не трогаем содержимое \\text{...}, \\mathrm{...}, \\operatorname{...}
        и других text/operator-арументов (там скобки — это сырой текст);
      - не трогаем экранированные \\( и \\).
    """
    def _mask_text_args(block: str) -> tuple[str, list[str]]:
        """Заменяет тело \\text{...} и аналогов на плейсхолдеры со счётом скобок."""
        out: list[str] = []
        stash: list[str] = []
        i = 0
        n = len(block)
        while i < n:
            m = _TEXT_ARG_RE.match(block, i)
            if not m:
                out.append(block[i])
                i += 1
                continue
            head_end = m.end()  # позиция сразу после "{"
            depth = 1
            j = head_end
            while j < n and depth > 0:
                c = block[j]
                if c == "\\" and j + 1 < n:
                    j += 2
                    continue
                if c == "{":
                    depth += 1
                elif c == "}":
                    depth -= 1
                j += 1
            inner = block[head_end:j - 1] if depth == 0 else block[head_end:j]
            placeholder = f"\x00TXT{len(stash)}\x00"
            stash.append(inner)
            out.append(block[i:head_end])
            out.append(placeholder)
            if depth == 0:
                out.append("}")
            i = j
        return "".join(out), stash

    def _rewrite(match: "re.Match[str]") -> str:
        block = match.group(0)
        block, stash = _mask_text_args(block)
        # Маркируем уже-завёрнутые \left( и \right), чтобы не дублировать.
        # Также маркируем экранированные \( и \) — это символ скобки в LaTeX.
        block = (
            block.replace(r"\left(", "\x00LP\x00")
                 .replace(r"\right)", "\x00RP\x00")
                 .replace(r"\(", "\x00ELP\x00")
                 .replace(r"\)", "\x00ERP\x00")
        )
        block = block.replace("(", r"\left(").replace(")", r"\right)")
        block = (
            block.replace("\x00LP\x00", r"\left(")
                 .replace("\x00RP\x00", r"\right)")
                 .replace("\x00ELP\x00", r"\(")
                 .replace("\x00ERP\x00", r"\)")
        )
        for idx, inner in enumerate(stash):
            block = block.replace(f"\x00TXT{idx}\x00", inner)
        return block

    return _MATH_BLOCK_RE.sub(_rewrite, content)


def _preprocess_html(content: str) -> str:
    content = _fix_math_parens(content)
    stripped = content.strip()
    low = stripped.lower()
    if low.startswith("<!doctype") or low.startswith("<html"):
        return content
    return (
        "<!doctype html>\n<html><head><meta charset=\"utf-8\"></head>"
        "<body>\n" + content + "\n</body></html>"
    )





_STYLE_TAG_RE = re.compile(
    r"<(?P<tag>span|p|div|h[1-6]|td|th|li)([^>]*?)style\s*=\s*(?:\"([^\"]+)\"|'([^']+)')([^>]*)>",
    flags=re.IGNORECASE,
)


def _extract_styled_spans(html: str) -> tuple[str, list[dict]]:
    """Помечает inline-styled элементы уникальными data-style-id для post-process.

    Возвращает (modified_html, list_of_styles).
    Каждый стиль: {"id": "STY0001", "props": {"color": "#c00000", ...}}.
    """
    styles: list[dict] = []

    def _parse_css(css: str) -> dict:
        out: dict = {}
        for decl in css.split(";"):
            decl = decl.strip()
            if not decl or ":" not in decl:
                continue
            k, _, v = decl.partition(":")
            out[k.strip().lower()] = v.strip()
        return out

    def _replace(m: re.Match) -> str:
        tag = m.group("tag")
        pre = m.group(2)
        css = m.group(3) if m.group(3) is not None else m.group(4)
        post = m.group(5)
        props = _parse_css(css)
        if not props:
            return m.group(0)
        sid = f"STY{len(styles):04d}"
        styles.append({"id": sid, "props": props, "tag": tag.lower()})
        return f'<{tag}{pre}data-sty-id="{sid}" style="{css}"{post}>'

    out = _STYLE_TAG_RE.sub(_replace, html)
    return out, styles


def _apply_styles_to_docx(docx_path, styles: list[dict]) -> None:
    """Post-process: применяет color/font-family/font-size/background к runs.

    Стратегия: проходим по всем параграфам/таблицам, для каждого run смотрим
    есть ли у него inline-маркер data-sty-id (pandoc HTML reader сохраняет
    кастомные атрибуты в data-... как часть class или просто игнорирует).
    Если pandoc не пробросил атрибут — fallback: ищем run, в котором текст
    совпадает по позиции с original span (грубое сопоставление).

    Для надёжности: pandoc bracketed_spans даёт {.cls #id key=val} → атрибуты
    custom уходят в Word custom properties, недоступных через python-docx.
    Поэтому делаем простой fallback — fuzzy matching по тексту между HTML
    span'ами и run'ами с тем же контентом.
    """
    if not styles:
        return

    from docx import Document
    from docx.shared import Pt, RGBColor

    def _parse_color(v: str) -> RGBColor | None:
        v = v.strip().lstrip("#")
        if len(v) == 3:
            v = "".join(c * 2 for c in v)
        if len(v) != 6:
            return None
        try:
            return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
        except ValueError:
            return None

    def _parse_size(v: str) -> Pt | None:
        v = v.strip().lower()
        try:
            if v.endswith("pt"):
                return Pt(float(v[:-2]))
            if v.endswith("px"):
                return Pt(float(v[:-2]) * _PX_TO_PT)
            if v.endswith("em"):
                return Pt(float(v[:-2]) * _EM_TO_PT)
            return Pt(float(v))
        except ValueError:
            return None

    doc = Document(str(docx_path))

    def _iter_paragraphs():
        yield from doc.paragraphs
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    paras = list(_iter_paragraphs())
    para_texts = [p.text for p in paras]

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }

    for style in styles:
        props = style["props"]
        tag = style.get("tag", "")
        snippet = style.get("text", "").strip()
        if not snippet:
            continue
        color = props.get("color")
        font_family = props.get("font-family")
        font_size = props.get("font-size")
        bg_color = props.get("background-color")
        text_align = props.get("text-align")
        rgb = _parse_color(color) if color else None
        bg_rgb = _parse_color(bg_color) if bg_color else None
        size_pt = _parse_size(font_size) if font_size else None
        font_name = font_family.split(",")[0].strip().strip("'\"") if font_family else None
        align_val = _align_map.get((text_align or "").lower()) if text_align else None

        # Сопоставление по `snippet in ptext` позиционно неточное: если один и
        # тот же текст повторяется в нескольких абзацах, стиль применится к
        # каждому совпадению. Отслеживать потреблённый offset здесь нельзя —
        # styles идут не в порядке появления в документе, а _apply_run_format
        # уже ищет первое вхождение внутри абзаца. Для inline-стилей это
        # приемлемо (повторы с разным стилем — редкий случай); точное
        # сопоставление потребовало бы проброса data-sty-id через pandoc.
        for p, ptext in zip(paras, para_texts):
            if snippet not in ptext:
                continue
            if rgb or font_name or size_pt or bg_rgb:
                _apply_run_format(p, snippet, rgb=rgb, font_name=font_name, size=size_pt, bg=bg_rgb)
            if align_val is not None and tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "div"):
                # Применяем выравнивание только к параграфам с paragraph-level тегами,
                # не к span/td/li (span не имеет смысла, td/li управляются отдельно).
                p.paragraph_format.alignment = align_val

    doc.save(str(docx_path))


def _apply_run_format(paragraph, snippet: str, *, rgb=None, font_name=None, size=None, bg=None) -> None:
    """Находит в paragraph runs, покрывающие snippet, и применяет форматирование.

    Контракт: НЕ трогает чужие элементы параграфа (math/hyperlinks/images)
    и НЕ трогает runs, не пересекающиеся со snippet. Затронутые runs
    разбиваются на before/styled/after; для каждого сегмента клонируется
    исходный run (сохраняя его текущее rPr — важно при многоступенчатом
    применении стилей), и к styled-клону доклеивается формат.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    full_text = paragraph.text
    pos = full_text.find(snippet)
    if pos < 0:
        return
    end = pos + len(snippet)

    # Для каждого затронутого run — список сегментов (text, is_styled).
    affected: list[tuple] = []  # (run_element, segments: list[(text, bool)])
    cursor = 0
    for r in list(paragraph.runs):
        rt = r.text or ""
        r_start = cursor
        r_end = cursor + len(rt)
        cursor = r_end
        if not rt:
            continue
        if r_end <= pos or r_start >= end:
            continue
        before = rt[: max(0, pos - r_start)] if r_start < pos else ""
        styled_start = max(0, pos - r_start)
        styled_end = min(len(rt), end - r_start)
        styled = rt[styled_start:styled_end]
        after = rt[styled_end:] if r_end > end else ""
        segs: list[tuple[str, bool]] = []
        if before:
            segs.append((before, False))
        if styled:
            segs.append((styled, True))
        if after:
            segs.append((after, False))
        if segs:
            affected.append((r._element, segs))

    if not affected or not any(is_st for _, segs in affected for _, is_st in segs):
        return

    def _set_text(run_elem, new_text: str) -> None:
        for t in run_elem.findall(qn("w:t")):
            run_elem.remove(t)
        # Удаляем не только табы, но и переносы строк (w:br/w:cr) — иначе
        # при перестроении текста run'а старые разрывы оставались бы и
        # дублировали/искажали содержимое.
        for tag in ("w:tab", "w:br", "w:cr"):
            for el in run_elem.findall(qn(tag)):
                run_elem.remove(el)
        t_el = OxmlElement("w:t")
        if new_text != new_text.strip():
            t_el.set(qn("xml:space"), "preserve")
        t_el.text = new_text
        run_elem.append(t_el)

    def _apply_style(run_elem) -> None:
        rpr = run_elem.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_elem.insert(0, rpr)
        if rgb is not None:
            for old in rpr.findall(qn("w:color")):
                rpr.remove(old)
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2]))
            rpr.append(color_el)
        if font_name:
            for old in rpr.findall(qn("w:rFonts")):
                rpr.remove(old)
            rfonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rfonts.set(qn(attr), font_name)
            rpr.append(rfonts)
        if size is not None:
            for old in rpr.findall(qn("w:sz")):
                rpr.remove(old)
            sz_el = OxmlElement("w:sz")
            sz_el.set(qn("w:val"), str(int(size.pt * 2)))
            rpr.append(sz_el)
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn("w:val"), str(int(size.pt * 2)))
            rpr.append(sz_cs)
        if bg is not None:
            for old in rpr.findall(qn("w:shd")):
                rpr.remove(old)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(bg[0], bg[1], bg[2]))
            rpr.append(shd)

    # Для каждого затронутого run: клонируем его, делаем сегменты,
    # вставляем последовательность клонов НА МЕСТО оригинала.
    for orig_elem, segs in affected:
        parent = orig_elem.getparent()
        if parent is None:
            continue
        anchor = orig_elem
        for text, is_styled in segs:
            clone = deepcopy(orig_elem)
            _set_text(clone, text)
            if is_styled:
                _apply_style(clone)
            anchor.addnext(clone)
            anchor = clone
        parent.remove(orig_elem)


_INNER_TAGS_RE = re.compile(r"<[^>]+>")


def _bind_styles_text(html_with_markers: str, styles: list[dict]) -> None:
    """Извлекает текст из помеченных span'ов и кладёт в styles[i]['text'].

    Ищем каждый стиль ПО ID независимо, чтобы вложенные span'ы внутри
    стилизованных p/div тоже корректно привязывались (re.finditer не
    возвращается внутрь уже сматченного блока).
    """
    for s in styles:
        sid = s["id"]
        pattern = re.compile(
            rf'<(?P<tag>span|p|div|h[1-6]|td|th|li)\b[^>]*data-sty-id="{sid}"[^>]*>(?P<body>.*?)</(?P=tag)>',
            flags=re.IGNORECASE | re.DOTALL,
        )
        m = pattern.search(html_with_markers)
        if not m:
            continue
        inner = _INNER_TAGS_RE.sub("", m.group("body"))
        s["text"] = inner.strip()


_TABLE_WIDTH_RE = re.compile(
    r"<table\b([^>]*)>", flags=re.IGNORECASE,
)
_WIDTH_IN_STYLE_RE = re.compile(
    r"width\s*:\s*([0-9]+(?:\.[0-9]+)?)\s*%", flags=re.IGNORECASE,
)


def _capture_table_widths(html: str) -> tuple[str, list[int | None]]:
    """Снимает ширину каждой таблицы (в % из style="width:N%") по порядку.

    Возвращает (html, [pct|None, ...]) — i-й элемент = ширина i-й таблицы в
    процентах (5000 = 100% в OOXML pct-единицах) либо None если не задана.
    Pandoc игнорирует width на таблице (ставит tblW=auto), поэтому захватываем
    сами и применяем в _ensure_table_borders, чтобы round-trip сохранял
    93%/82% и т.п. вместо форса 100%.
    """
    widths: list[int | None] = []
    for m in _TABLE_WIDTH_RE.finditer(html):
        attrs = m.group(1) or ""
        wm = _WIDTH_IN_STYLE_RE.search(attrs)
        if wm:
            pct = round(float(wm.group(1)) * 50)  # 100% → 5000
            widths.append(max(1, min(pct, 5000)))
        else:
            widths.append(None)
    return html, widths


def _post_process_docx(docx_path, table_widths: list | None = None) -> None:
    """Добавляет границы всем таблицам и page-break перед каждым H1.

    Pandoc применяет кастомный стиль 'Table' из reference, но Word/LibreOffice
    могут проигнорировать табличные границы из стиля. Дублируем их inline
    в каждую таблицу.
    """
    from docx import Document

    doc = Document(str(docx_path))

    widths = table_widths or []
    for i, tbl in enumerate(doc.tables):
        pct = widths[i] if i < len(widths) else None
        _ensure_table_borders(tbl, width_pct=pct)

    h1_seen = False
    for para in doc.paragraphs:
        if (para.style.name or "" if para.style else "").lower() == "heading 1":
            if h1_seen:
                _insert_page_break_before(para)
            h1_seen = True

    _fix_paragraph_indents(doc)

    doc.save(str(docx_path))


def _para_has_image(paragraph) -> bool:
    """True, если параграф содержит хотя бы одну картинку (w:drawing/pict)."""
    from docx.oxml.ns import qn
    el = paragraph._element
    return bool(el.findall(".//" + qn("w:drawing")) or el.findall(".//" + qn("w:pict")))


def _para_alignment(paragraph) -> str | None:
    """Возвращает 'center'/'right'/'justify'/'left' либо None, читая pPr/jc напрямую."""
    from docx.oxml.ns import qn
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        return None
    return (jc.get(qn("w:val")) or "").lower() or None


def _zero_first_line_indent(paragraph) -> None:
    """Обнуляет red-line (firstLine) у параграфа, сохраняя прочие отступы."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # Снимаем любой положительный firstLine, не вводим отрицательный hanging.
    ind.set(qn("w:firstLine"), "0")
    if ind.get(qn("w:hanging")) is not None:
        del ind.attrib[qn("w:hanging")]


def _fix_paragraph_indents(doc) -> None:
    """Убирает паразитный red-line 1.25 cm там, где он ломает компоновку.

    Pandoc вешает firstLine из стиля Normal на КАЖДЫЙ <p>, включая:
      • центрированные/right-выровненные подписи и картинки — отступ сдвигает
        первую строку «Рисунок N» вправо (видно когда подпись в две строки);
      • первый текстовый абзац сразу ПОСЛЕ картинки — выглядит как сдвиг без
        причины.
    Заголовки (Heading*) и ячейки таблиц обрабатываются отдельно — не трогаем.
    """
    prev_was_image = False
    for para in doc.paragraphs:
        style_name = (para.style.name or "" if para.style else "").lower()
        is_heading = style_name.startswith("heading") or style_name in ("title", "subtitle")
        has_img = _para_has_image(para)
        align = _para_alignment(para)

        if not is_heading:
            if align in ("center", "right"):
                _zero_first_line_indent(para)
            elif prev_was_image and para.text.strip():
                _zero_first_line_indent(para)

        prev_was_image = has_img


def _ensure_table_borders(table, width_pct: int | None = None) -> None:
    """Добавляет границы + фиксирует ширину таблицы (100%) с равными колонками.

    Pandoc ставит tblW=auto/0, из-за чего LibreOffice не понимает как
    отрисовать таблицу и схлопывает её. Принудительно ставим tblW=pct/5000
    (=100%) и равные ширины колонок.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_el = table._element
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)

    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    # Уважаем ширину из HTML (style="width:N%"); по умолчанию — 100%.
    tblW.set(qn("w:w"), str(width_pct if width_pct is not None else 5000))
    tblPr.append(tblW)

    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    n_cols = len(table.columns)
    if n_cols > 0:
        total_pct = width_pct if width_pct is not None else 5000
        col_width = total_pct // n_cols
        tblGrid = tbl_el.find(qn("w:tblGrid"))
        if tblGrid is not None:
            tbl_el.remove(tblGrid)
        tblGrid = OxmlElement("w:tblGrid")
        for _ in range(n_cols):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(_PAGE_TEXT_WIDTH_TWIPS // n_cols))
            tblGrid.append(gc)
        tbl_el.insert(list(tbl_el).index(tblPr) + 1, tblGrid)

        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:type"), "pct")
                tcW.set(qn("w:w"), str(col_width))
                tcPr.append(tcW)

                # Сбрасываем стили параграфов в ячейках: убираем
                # FirstParagraph/BodyText (они дают w:firstLine из Normal),
                # устанавливаем пустой стиль = inherited Normal с обнулённым отступом.
                for p in cell.paragraphs:
                    p_el = p._element
                    pPr = p_el.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p_el.insert(0, pPr)
                    for old in pPr.findall(qn("w:pStyle")):
                        pPr.remove(old)
                    for old in pPr.findall(qn("w:ind")):
                        pPr.remove(old)
                    ind = OxmlElement("w:ind")
                    ind.set(qn("w:firstLine"), "0")
                    ind.set(qn("w:left"), "0")
                    pPr.append(ind)
                    for old in pPr.findall(qn("w:spacing")):
                        pPr.remove(old)
                    sp = OxmlElement("w:spacing")
                    sp.set(qn("w:line"), "240")
                    sp.set(qn("w:lineRule"), "auto")
                    sp.set(qn("w:before"), "0")
                    sp.set(qn("w:after"), "0")
                    pPr.append(sp)


def _insert_page_break_before(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    for old in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(old)
    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)


def create_docx(call: ToolCall) -> ToolResult:
    """HTML → DOCX через pandoc + post-processing inline-стилей.

    args:
      path (str, required) — путь к выходному .docx
      content (str, required) — HTML-разметка (LaTeX $...$, $$...$$ поддерживается)
      reference_doc (str, optional) — путь к шаблону стилей (.docx).
          Если не указан — используется дефолтный (TNR 14, 1.5, чёрные заголовки).
          Передай "none" чтобы отключить дефолтный шаблон.
      overwrite (bool, optional, default=True) — перезаписать существующий файл
    """
    op_id = uuid.uuid4().hex[:8]
    args = call.args or {}

    path_str = clean_path(args.get("path", ""))
    if not path_str:
        return ToolResult(
            name="create_docx", status="error",
            output="Output file path (path) not specified.",
            exit_code=1, command=call.command,
        )

    content = args.get("content", "")
    if content is None:
        content = ""
    if not isinstance(content, str):
        content = str(content)
    if not content.strip():
        return ToolResult(
            name="create_docx", status="error",
            output="Empty HTML content (content).",
            exit_code=1, command=call.command,
        )

    out_path = resolve_path(path_str)
    if out_path.suffix.lower() != ".docx":
        return ToolResult(
            name="create_docx", status="error",
            output=f"Path must end with .docx, got: {path_str}",
            exit_code=1, command=call.command,
        )

    overwrite = bool(args.get("overwrite", True))
    if out_path.exists() and not overwrite:
        return ToolResult(
            name="create_docx", status="error",
            output=f"File already exists: {path_str} (overwrite=false).",
            exit_code=1, command=call.command, fatal=True,
        )

    pandoc = _find_pandoc()
    if not pandoc:
        logger.warning("create_docx[{}]: pandoc not in PATH", op_id)
        return ToolResult(
            name="create_docx", status="error",
            output=_install_hint(),
            exit_code=127, command=call.command, fatal=True,
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)

    content = _strip_read_marker(content)
    # Снимок чистого HTML ДО code/stray-стэша: по нему WRITE-side восстановит
    # точный whitespace (pandoc html→docx снова обрежет ведущие/внутренние
    # пробелы и табы). Сопоставление идёт по схлопнутому тексту, поэтому нужен
    # именно исходный HTML без буквенных плейсхолдеров.
    _ws_source_html = content
    # Сохраняем чистый HTML-исходник (после strip-маркера, до pandoc-препроцесса)
    # для последующего редактирования через patch_file / точного round-trip чтения.
    _save_docx_source(out_path, content)

    # ── Точный round-trip через шаблон-оригинал ──
    # Если этот HTML получен из ранее прочитанного внешнего docx (есть шаблон) и
    # текст совпадает — клонируем оригинал и патчим только текст. Так сохраняются
    # секции, стили, колонтитулы, нумерация, картинки, ширины таблиц и вёрстка
    # (число страниц) — то, что pandoc безвозвратно теряет. При плохом совпадении
    # (тяжёлая перезапись) функция вернёт False и мы откатимся на pandoc.
    reference_arg = args.get("reference_doc")
    if reference_arg is None or (
        isinstance(reference_arg, str) and reference_arg.strip().lower() != "none"
    ):
        try:
            template = _find_matching_template(out_path, content)
            if template is not None:
                from tools.file_ops._docx_whitespace import write_via_template
                if write_via_template(template, out_path, content):
                    logger.info("create_docx[{}]: exact round-trip via template {}", op_id, template.name)
                    return ToolResult(
                        name="create_docx", status="ok",
                        output=f"✓ DOCX created (exact template round-trip): {path_str}",
                        exit_code=0, command=call.command,
                    )
        except Exception as e:
            logger.opt(exception=True).warning(
                "create_docx[{}]: template path failed, falling back to pandoc: {}", op_id, e,
            )
    # Выносим тело code-блоков в стэш ПЕРВЫМ шагом и ставим буквенные
    # плейсхолдеры. pandoc HTML reader безвозвратно съедает <iostream>,
    # vector<Lexeme>, <программа> внутри <code> (декодирует сущности и
    # выкидывает tag-подобные токены), поэтому код вообще не пропускаем через
    # pandoc — впечатываем сырой текст в docx после конвертации.
    content, _code_stash = _extract_code_blocks(content)
    # ВАЖНО: НЕ схлопываем пустые параграфы. Это ломало round-trip — read
    # отдаёт ровно столько <p><br/></p>, сколько было в исходном docx, а
    # схлопывание до 2 рушило вёрстку титульника (текст съезжал, таблицы-рамки
    # уезжали вниз). Правило «макс 2-3 пустых на титульнике» остаётся в навыке
    # docx-mastery как рекомендация модели при создании документа с нуля.
    content, _table_widths = _capture_table_widths(content)
    content = _wrap_table_cells(content)
    # ВАЖНО: extract_stray_angles ПОСЛЕ wrap_table_cells — _CellWrapper гоняет
    # HTML через HTMLParser, который декодирует сущности. Выносим литеральные
    # <...> (BNF/плейсхолдеры) в стэш и впечатываем обратно после конвертации
    # (restore_stray_angles), т.к. pandoc их съедает как мнимые теги.
    content, _stray_stash = _extract_stray_angles(content)
    content, styles = _extract_styled_spans(content)
    _bind_styles_text(content, styles)
    content = _preprocess_html(content)

    # reference_arg уже считан выше (перед шаблонным путём).
    ref_path = None
    if reference_arg is None:
        try:
            ref_path = get_default_reference_path()
        except Exception as e:
            logger.opt(exception=True).warning(
                "create_docx[{}]: cannot build default reference.docx: {}", op_id, e,
            )
            ref_path = None
    elif isinstance(reference_arg, str) and reference_arg.strip().lower() == "none":
        ref_path = None
    else:
        cand = resolve_path(clean_path(str(reference_arg)))
        if not cand.exists():
            return ToolResult(
                name="create_docx", status="error",
                output=f"reference_doc not found: {reference_arg}",
                exit_code=1, command=call.command,
            )
        ref_path = cand

    cmd = [
        pandoc,
        "-f", "html+tex_math_dollars+raw_html",
        "-t", "docx",
        "-o", str(out_path),
        "--standalone",
        "--wrap=preserve",
    ]
    if ref_path is not None:
        cmd += [f"--reference-doc={ref_path}"]

    logger.info(
        "create_docx[{}]: pandoc → {} (html={}b, ref={}, styled_spans={})",
        op_id, path_str, len(content.encode("utf-8")),
        ref_path.name if ref_path else "none", len(styles),
    )

    try:
        proc = subprocess.run(
            cmd,
            input=content,
            text=True,
            capture_output=True,
            cwd=get_working_dir(),
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        logger.error("create_docx[{}]: pandoc timeout (120s)", op_id)
        return ToolResult(
            name="create_docx", status="error",
            output="Pandoc did not respond within 120 seconds — conversion aborted.",
            exit_code=124, command=call.command,
        )
    except Exception as e:
        logger.opt(exception=True).error(
            "create_docx[{}]: pandoc spawn failed: {}", op_id, e,
        )
        return ToolResult(
            name="create_docx", status="error",
            output=f"Failed to start pandoc: {type(e).__name__}: {e}",
            exit_code=1, command=call.command,
        )

    if proc.returncode != 0:
        stderr = (proc.stderr or "").strip()[:1500]
        logger.warning(
            "create_docx[{}]: pandoc exit={} stderr={!r}",
            op_id, proc.returncode, stderr[:300],
        )
        return ToolResult(
            name="create_docx", status="error",
            output=f"Pandoc exited with code {proc.returncode}:\n{stderr}",
            exit_code=proc.returncode, command=call.command,
        )

    if not out_path.exists():
        return ToolResult(
            name="create_docx", status="error",
            output="Pandoc finished without errors, but output file was not created.",
            exit_code=1, command=call.command,
        )

    # ВАЖНО: restore_stray_angles ДО apply_styles — _apply_styles_to_docx
    # сопоставляет snippet (с реальными <…>) с текстом параграфа. Если в docx
    # ещё стоит плейсхолдер NECLISTRAYANGLE…, совпадения нет и стиль (напр.
    # font-size:8pt для строк кода) не применяется. Сначала впечатываем скобки.
    if _stray_stash:
        try:
            _restore_stray_angles(out_path, _stray_stash)
            logger.info("create_docx[{}]: restored {} stray angle(s)", op_id, len(_stray_stash))
        except Exception as e:
            logger.opt(exception=True).warning(
                "create_docx[{}]: stray-angle restore failed: {}", op_id, e,
            )

    if styles:
        try:
            _apply_styles_to_docx(out_path, styles)
            logger.info("create_docx[{}]: applied {} inline styles", op_id, len(styles))
        except Exception as e:
            logger.opt(exception=True).warning(
                "create_docx[{}]: post-style failed: {}", op_id, e,
            )

    try:
        _post_process_docx(out_path, table_widths=_table_widths)
        logger.info("create_docx[{}]: post-process ok (table borders, h1 page breaks)", op_id)
    except Exception as e:
        logger.opt(exception=True).warning(
            "create_docx[{}]: post-process failed: {}", op_id, e,
        )

    if _code_stash:
        try:
            _restore_code_blocks(out_path, _code_stash)
            logger.info("create_docx[{}]: restored {} code block(s)", op_id, len(_code_stash))
        except Exception as e:
            logger.opt(exception=True).warning(
                "create_docx[{}]: code restore failed: {}", op_id, e,
            )

    # Восстанавливаем точный whitespace, который pandoc html→docx обрезал
    # (ведущие/хвостовые/внутренние пробелы и табы). Запускаем ПОСЛЕДНИМ —
    # текст параграфов уже окончательный (code/stray впечатаны), сопоставление
    # по схлопнутому тексту однозначно.
    try:
        from tools.file_ops._docx_whitespace import restore_into_docx as _ws_restore_docx
        _ws_n = _ws_restore_docx(out_path, _ws_source_html)
        if _ws_n:
            logger.info("create_docx[{}]: restored whitespace in {} paragraph(s)", op_id, _ws_n)
    except Exception as e:
        logger.opt(exception=True).warning(
            "create_docx[{}]: whitespace restore failed: {}", op_id, e,
        )

    size = out_path.stat().st_size
    stderr_tail = (proc.stderr or "").strip()
    warn = f"\n⚠ pandoc stderr: {stderr_tail[:400]}" if stderr_tail else ""
    logger.info("create_docx[{}]: ok {} ({}b)", op_id, path_str, size)
    return ToolResult(
        name="create_docx", status="ok",
        output=f"✓ DOCX created: {path_str}{warn}",
        exit_code=0, command=call.command,
    )


__all__ = ["create_docx"]