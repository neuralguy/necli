"""Сохранение точного whitespace при docx round-trip (read ↔ write).

Проблема: pandoc на ОБОИХ направлениях обрезает ведущие/хвостовые пробелы и
схлопывает внутренние пробельные серии в один пробел. Реальные документы
(особенно с ручной разметкой пробелами для центрирования и с отступами кода)
из-за этого теряют ~14% символов на round-trip.

python-docx читает текст параграфа с ТОЧНЫМ whitespace (ведущие пробелы, \t,
внутренние серии, хвостовые). Мы используем это, чтобы восстановить пробелы:
  • на ЧТЕНИИ — впечатываем точный whitespace обратно в HTML (после pandoc);
  • на ЗАПИСИ — после pandoc html→docx впечатываем whitespace обратно в runs
    с xml:space=preserve, т.к. pandoc снова его обрезал.

Сопоставление параграфов docx ↔ HTML делаем через обход документа в реальном
порядке (body-параграфы + спуск в ячейки таблиц) и difflib.SequenceMatcher по
whitespace-СХЛОПНУТОМУ тексту — это устойчиво к перемежению таблиц и к
потерянным пустым параграфам (проверено: 502/507 блоков, ratio 0.978).
"""

from __future__ import annotations

import difflib
import re

from logger import logger


def collapse_ws(s: str) -> str:
    """Ключ сопоставления: схлопывает любые пробельные серии в один пробел и
    обрезает концы — ровно так, как pandoc нормализует текст параграфа."""
    return re.sub(r"\s+", " ", s).strip()


def has_significant_ws(s: str) -> bool:
    """True, если у непустого текста есть whitespace, который pandoc потеряет:
    ведущие/хвостовые пробелы/табы или внутренние серии 2+ пробелов/таб."""
    if not s.strip():
        return False
    return s != collapse_ws(s)


def tree_walk_paragraphs(doc) -> list:
    """Параграфы docx в РЕАЛЬНОМ порядке документа (как сериализует pandoc):
    body-level <p>, при встрече <tbl> — спускаемся в строки/ячейки/параграфы.

    Возвращает список объектов python-docx Paragraph (в т.ч. из ячеек).
    """
    from docx.oxml.ns import qn

    elem_to_para: dict[int, object] = {}
    for p in doc.paragraphs:
        elem_to_para[id(p._element)] = p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    elem_to_para[id(cp._element)] = cp

    ordered: list = []

    def _walk(parent) -> None:
        for child in parent.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                p = elem_to_para.get(id(child))
                if p is not None:
                    ordered.append(p)
            elif tag == "tbl":
                for trow in child.findall(qn("w:tr")):
                    for tc in trow.findall(qn("w:tc")):
                        _walk(tc)

    _walk(doc.element.body)
    return ordered


_BLOCK_RE = re.compile(
    r"<(?P<tag>p|h[1-6])\b(?P<attrs>[^>]*)>(?P<inner>.*?)</(?P=tag)>",
    flags=re.IGNORECASE | re.DOTALL,
)
# Ячейки <td>/<th>, чьё содержимое — голый текст (pandoc не оборачивает в <p>
# простые одно-абзацные ячейки). Многоабзацные ячейки ловятся _BLOCK_RE по <p>.
_CELL_RE = re.compile(
    r"<(?P<tag>td|th)\b(?P<attrs>[^>]*)>(?P<inner>.*?)</(?P=tag)>",
    flags=re.IGNORECASE | re.DOTALL,
)
_TAG_RE = re.compile(r"<[^>]+>")
_BLOCK_CHILD_RE = re.compile(r"<(?:p|h[1-6])\b", flags=re.IGNORECASE)


def parse_html_blocks(html: str) -> list[dict]:
    """Все <p>/<h*> блоки + голотекстовые <td>/<th> по порядку (offset в строке).

    Для каждого: span inner (start/end), inner-HTML и его text-content
    (теги убраны, сущности раскрыты) для сопоставления.
    """
    import html as _htmlmod

    blocks: list[dict] = []

    def _add(m, tag):
        inner = m.group("inner")
        blocks.append(
            {
                "tag": tag,
                "attrs": m.group("attrs") or "",
                "inner": inner,
                "inner_start": m.start("inner"),
                "inner_end": m.end("inner"),
                "text": _htmlmod.unescape(_TAG_RE.sub("", inner)),
            }
        )

    for m in _BLOCK_RE.finditer(html):
        _add(m, m.group("tag"))
    # Голотекстовые ячейки: только если внутри НЕТ <p>/<h*> (иначе уже учтено).
    for m in _CELL_RE.finditer(html):
        if _BLOCK_CHILD_RE.search(m.group("inner")):
            continue
        if not m.group("inner").strip():
            continue
        _add(m, m.group("tag"))

    # Восстанавливаем документный порядок по позиции inner.
    blocks.sort(key=lambda b: b["inner_start"])
    return blocks


def align_pairs(docx_texts: list[str], html_texts: list[str]) -> list[tuple[int, int]]:
    """Сопоставляет индексы docx↔html по whitespace-схлопнутому тексту.

    Возвращает только пары из equal-блоков SequenceMatcher (надёжные).
    """
    a = [collapse_ws(t) for t in docx_texts]
    b = [collapse_ws(t) for t in html_texts]
    sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
    pairs: list[tuple[int, int]] = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag != "equal":
            continue
        for k in range(i2 - i1):
            pairs.append((i1 + k, j1 + k))
    return pairs


_LEAD_WS_RE = re.compile(r"^\s+")
_TRAIL_WS_RE = re.compile(r"\s+$")
_HAS_TAG_RE = re.compile(r"<[^>]+>")


def _split_ws(text: str) -> tuple[str, str, str]:
    """(leading_ws, core, trailing_ws) для текста параграфа."""
    lead_m = _LEAD_WS_RE.match(text)
    lead = lead_m.group(0) if lead_m else ""
    trail_m = _TRAIL_WS_RE.search(text)
    trail = trail_m.group(0) if trail_m else ""
    core = text[len(lead): len(text) - len(trail)] if (lead or trail) else text
    return lead, core, trail


_SEG_RE = re.compile(r"<[^>]+>|[^<]+")


def _restore_ws_in_inner_html(inner: str, exact: str) -> str | None:
    """Раскладывает точный текст `exact` по ТЕКСТОВЫМ сегментам inner-HTML,
    не трогая теги. Возвращает новый inner или None, если не-пробельные символы
    не сошлись (тогда вызывающий откатится на lead/trail).

    Сегменты: либо тег `<...>` (пропускаем дословно), либо текст (его не-ws
    символы сохраняем, ws берём из exact). Сущности (&amp; и т.п.) раскрываем
    для подсчёта, на выходе пере-экранируем < > &.
    """
    import html as _h

    segs = _SEG_RE.findall(inner)
    # Текстовые сегменты (декодированные) — по ним раскладываем exact.
    text_seg_idx = [i for i, s in enumerate(segs) if not s.startswith("<")]
    text_pieces = [_h.unescape(segs[i]) for i in text_seg_idx]
    if collapse_ws("".join(text_pieces)) != collapse_ws(exact):
        return None
    new_pieces = _redistribute_exact(exact, text_pieces)
    if new_pieces is None:
        return None
    out_segs = list(segs)
    for slot, piece in zip(text_seg_idx, new_pieces):
        out_segs[slot] = _h.escape(piece, quote=False)
    return "".join(out_segs)


def restore_into_html(html: str, doc) -> str:
    """READ-side: впечатывает точный whitespace из docx обратно в HTML.

    pandoc уже обрезал пробелы в <p>; восстанавливаем их, сопоставив параграфы
    docx (tree-walk) с HTML-блоками (SequenceMatcher по схлопнутому тексту).

    Скоуп:
      • ведущий/хвостовой whitespace — впечатываем в начало/конец inner-HTML;
      • внутренние серии — только если блок без inline-тегов (чистый текст),
        тогда заменяем весь inner на точный текст docx.
    Параграфы без значимого whitespace не трогаем.
    """
    try:
        paras = tree_walk_paragraphs(doc)
        docx_texts = [p.text for p in paras]
        ne_idx = [i for i, t in enumerate(docx_texts) if t.strip()]
        ne_texts = [docx_texts[i] for i in ne_idx]

        blocks = parse_html_blocks(html)
        block_texts = [b["text"] for b in blocks]

        pairs = align_pairs(ne_texts, block_texts)

        # Готовим правки как (inner_start, inner_end, new_inner); применяем
        # справа налево, чтобы не сдвигать офсеты.
        edits: list[tuple[int, int, str]] = []
        for di, bj in pairs:
            docx_text = ne_texts[di]
            if not has_significant_ws(docx_text):
                continue
            blk = blocks[bj]
            inner = blk["inner"]
            has_tags = bool(_HAS_TAG_RE.search(inner))
            lead, _core, trail = _split_ws(docx_text)

            if not has_tags:
                # Чистый текст: заменяем весь inner на точный текст docx,
                # но если в HTML inner есть собственный whitespace/сущности —
                # сравниваем по схлопнутому, чтобы не затереть отличия pandoc.
                if collapse_ws(blk["text"]) == collapse_ws(docx_text):
                    import html as _h
                    new_inner = _h.escape(docx_text, quote=False)
                    if new_inner != inner:
                        edits.append((blk["inner_start"], blk["inner_end"], new_inner))
                continue

            # Есть inline-теги: раскладываем точный whitespace по текстовым
            # сегментам inner-HTML (части между тегами), не трогая сами теги.
            new_inner = _restore_ws_in_inner_html(inner, docx_text)
            if new_inner is None:
                # Не сошлось — восстановим хотя бы ведущий/хвостовой.
                new_inner = inner
                if lead and not _LEAD_WS_RE.match(new_inner):
                    new_inner = lead + new_inner
                if trail and not _TRAIL_WS_RE.search(new_inner):
                    new_inner = new_inner + trail
            if new_inner != inner:
                edits.append((blk["inner_start"], blk["inner_end"], new_inner))

        if not edits:
            return html
        edits.sort(key=lambda e: e[0], reverse=True)
        out = html
        for s, e, new in edits:
            out = out[:s] + new + out[e:]
        logger.debug("docx ws restore (read): {} paragraph(s) adjusted", len(edits))
        return out
    except Exception as exc:  # никогда не ломаем чтение
        logger.opt(exception=True).debug("docx ws restore (read) skipped: {}", exc)
        return html


def _set_para_leading_ws(paragraph, lead: str) -> None:
    """Впечатывает ведущий whitespace в ПЕРВЫЙ run параграфа (xml:space=preserve)."""
    from docx.oxml.ns import qn
    runs = paragraph.runs
    if not runs:
        return
    for r in runs:
        t_el = r._element.find(qn("w:t"))
        if t_el is None or t_el.text is None:
            continue
        if (t_el.text or "").startswith(lead):
            return  # уже на месте
        t_el.text = lead + t_el.text
        t_el.set(qn("xml:space"), "preserve")
        return


def _set_para_trailing_ws(paragraph, trail: str) -> None:
    from docx.oxml.ns import qn
    runs = paragraph.runs
    if not runs:
        return
    for r in reversed(runs):
        t_el = r._element.find(qn("w:t"))
        if t_el is None or t_el.text is None:
            continue
        if (t_el.text or "").endswith(trail):
            return
        t_el.text = t_el.text + trail
        t_el.set(qn("xml:space"), "preserve")
        return


def _redistribute_exact(exact: str, run_texts: list[str]) -> list[str] | None:
    """Раскладывает точный текст `exact` по run'ам, сохраняя у каждого run его
    НЕ-пробельные символы, а пробелы беря из `exact`.

    Предусловие: конкатенация run_texts схлопывается к тому же, что и exact
    (т.е. совпадают все не-пробельные символы по порядку). Иначе → None.

    Так восстанавливаются ведущие/внутренние/хвостовые пробелы и табы, даже
    когда текст разбит pandoc'ом на несколько runs (напр. '}' и ';' отдельно).
    """
    produced = "".join(run_texts)
    if collapse_ws(produced) != collapse_ws(exact):
        return None
    res: list[str] = []
    ti = 0
    tn = len(exact)
    last = len(run_texts) - 1
    for ridx, rt in enumerate(run_texts):
        nonws_count = sum(1 for c in rt if not c.isspace())
        buf: list[str] = []
        got = 0
        while ti < tn:
            c = exact[ti]
            if c.isspace():
                buf.append(c)
                ti += 1
            elif got < nonws_count:
                buf.append(c)
                got += 1
                ti += 1
            else:
                break
        res.append("".join(buf))
    # Остаток (хвостовые пробелы) — в последний run.
    if ti < tn:
        res[last] += exact[ti:]
    return res


def _restore_exact_text_across_runs(paragraph, exact: str) -> bool:
    """Переписывает текстовые runs параграфа так, чтобы их конкатенация была
    ровно `exact` (с preserve). Возвращает True, если применено."""
    from docx.oxml.ns import qn
    text_els = [
        r._element.find(qn("w:t"))
        for r in paragraph.runs
        if r._element.find(qn("w:t")) is not None
    ]
    if not text_els:
        return False
    run_texts = [(te.text or "") for te in text_els]
    if "".join(run_texts) == exact:
        return True
    new_texts = _redistribute_exact(exact, run_texts)
    if new_texts is None:
        return False
    for te, new in zip(text_els, new_texts):
        if te.text != new:
            te.text = new
            te.set(qn("xml:space"), "preserve")
    return True


def restore_into_docx(docx_path, source_html: str) -> int:
    """WRITE-side: впечатывает точный whitespace из исходного HTML в готовый docx.

    pandoc html→docx снова обрезал пробелы; восстанавливаем их, сопоставив
    HTML-блоки (исходный текст до pandoc) с параграфами получившегося docx
    (tree-walk) через SequenceMatcher по схлопнутому тексту.

    Возвращает число изменённых параграфов. Никогда не бросает наружу.
    """
    try:
        from docx import Document

        blocks = parse_html_blocks(source_html)
        html_texts = [b["text"] for b in blocks]
        # Блоки HTML со значимым whitespace — только их нужно восстанавливать.
        ne_html_idx = [i for i, t in enumerate(html_texts) if t.strip()]
        ne_html_texts = [html_texts[i] for i in ne_html_idx]

        doc = Document(str(docx_path))
        paras = tree_walk_paragraphs(doc)
        para_texts = [p.text for p in paras]
        ne_para_idx = [i for i, t in enumerate(para_texts) if t.strip()]
        ne_para_texts = [para_texts[i] for i in ne_para_idx]

        # align: html (a) ↔ docx-paras (b)
        pairs = align_pairs(ne_html_texts, ne_para_texts)

        changed = 0
        for hi, pi in pairs:
            html_text = ne_html_texts[hi]
            if not has_significant_ws(html_text):
                continue
            para = paras[ne_para_idx[pi]]
            # Общий путь: раскладываем точный текст по runs (ведущие/внутренние/
            # хвостовые пробелы и табы сразу). Работает и при разбиении на runs.
            if _restore_exact_text_across_runs(para, html_text):
                changed += 1
                continue
            # Fallback: не сошлись не-пробельные символы (редко: pandoc заменил
            # символ) — восстанавливаем хотя бы ведущий/хвостовой whitespace.
            lead, _core, trail = _split_ws(html_text)
            did = False
            if lead:
                _set_para_leading_ws(para, lead)
                did = True
            if trail:
                _set_para_trailing_ws(para, trail)
                did = True
            if did:
                changed += 1

        # Пустые параграфы: read отдаёт <p><br/></p> (иначе pandoc выкинет
        # пустой абзац вовсе), но <br/> превращается в run с <w:br/> = текст
        # '\n'. В оригинале пустые абзацы — без runs. Убираем одиночный <w:br/>
        # из абзацев, где он единственное содержимое, чтобы пустые абзацы были
        # байт-идентичны оригиналу.
        cleaned = _strip_lone_breaks(doc)

        if changed or cleaned:
            doc.save(str(docx_path))
        return changed
    except Exception as exc:  # никогда не ломаем запись
        logger.opt(exception=True).warning("docx ws restore (write) skipped: {}", exc)
        return 0


def _strip_lone_breaks(doc) -> int:
    """Удаляет run с единственным <w:br/> из иначе-пустых параграфов.

    Возвращает число очищенных параграфов. Так пустой абзац <p><br/></p>,
    нужный pandoc'у для сохранения абзаца, после записи становится по-настоящему
    пустым (как в исходном docx), без паразитного '\n'.
    """
    from docx.oxml.ns import qn

    br = qn("w:br")
    rpr = qn("w:rPr")
    n = 0
    for para in tree_walk_paragraphs(doc):
        if para.text.strip():
            continue  # есть видимый текст — не трогаем
        runs = para.runs
        if not runs:
            continue  # уже по-настоящему пустой
        # Содержимое каждого run (без rPr) должно быть только <w:br/> или пусто.
        def _content(r):
            return [c.tag for c in r._element if c.tag != rpr]
        if not all(set(_content(r)) <= {br} for r in runs):
            continue
        # Убираем такие runs — абзац становится по-настоящему пустым.
        removed = False
        for r in list(runs):
            r._element.getparent().remove(r._element)
            removed = True
        if removed:
            n += 1
    return n


def _set_paragraph_text_exact(paragraph, exact: str) -> None:
    """Принудительно делает текст параграфа равным `exact`, сохраняя формат
    ПЕРВОГО текстового run (шрифт/размер/цвет). Остальные текстовые runs
    очищаются. Используется для ОТРЕДАКТИРОВАННЫХ абзацев в шаблонном пути,
    где не-пробельные символы не совпали с оригиналом.
    """
    from docx.oxml.ns import qn
    text_runs = [r for r in paragraph.runs if r._element.find(qn("w:t")) is not None]
    if not text_runs:
        if paragraph.runs:
            paragraph.runs[0].text = exact
        else:
            paragraph.add_run(exact)
        return
    first = text_runs[0]
    t_el = first._element.find(qn("w:t"))
    t_el.text = exact
    t_el.set(qn("xml:space"), "preserve")
    for r in text_runs[1:]:
        te = r._element.find(qn("w:t"))
        if te is not None:
            te.text = ""


def write_via_template(template_path, output_path, desired_html: str) -> bool:
    """Точный round-trip: клонирует оригинал-шаблон и патчит ТОЛЬКО текст
    абзацев под содержимое desired_html. Секции, стили, колонтитулы, нумерация,
    картинки, ширины таблиц остаются от оригинала (pandoc их теряет).

    Для немодифицированного чтения→записи результат байт-идентичен оригиналу.
    При лёгких правках патчатся совпавшие по позиции абзацы.

    Возвращает True при успехе. Если выравнивание плохое (тяжёлые
    структурные правки) — возвращает False, и вызывающий откатывается на pandoc.
    """
    try:
        from docx import Document

        blocks = parse_html_blocks(desired_html)
        html_texts = [b["text"] for b in blocks]
        ne_html = [t for t in html_texts if t.strip()]

        tpl_doc = Document(str(template_path))
        tpl_paras = tree_walk_paragraphs(tpl_doc)
        tpl_texts = [p.text for p in tpl_paras]
        ne_tpl_idx = [i for i, t in enumerate(tpl_texts) if t.strip()]
        ne_tpl_texts = [tpl_texts[i] for i in ne_tpl_idx]

        # Выравнивание: html (a) ↔ template-paras (b)
        a = [collapse_ws(t) for t in ne_html]
        b = [collapse_ws(t) for t in ne_tpl_texts]
        sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
        opcodes = sm.get_opcodes()

        equal = sum((i2 - i1) for tag, i1, i2, j1, j2 in opcodes if tag == "equal")
        total = max(len(a), 1)
        coverage = equal / total
        # Если совпало мало — это не round-trip, а серьёзная перезапись:
        # шаблонный путь не подходит, пусть работает pandoc.
        if coverage < 0.85:
            logger.info(
                "docx template path declined: coverage={:.2f} (html={} tpl={})",
                coverage, len(a), len(b),
            )
            return False

        # Патчим текст совпавших абзацев точным html-текстом.
        patched = 0
        for tag, i1, i2, j1, j2 in opcodes:
            if tag not in ("equal", "replace"):
                continue
            # сопоставляем поэлементно в пределах блока
            for k in range(min(i2 - i1, j2 - j1)):
                want = ne_html[i1 + k]
                para = tpl_paras[ne_tpl_idx[j1 + k]]
                if para.text == want:
                    continue
                if not _restore_exact_text_across_runs(para, want):
                    _set_paragraph_text_exact(para, want)
                patched += 1

        # Гарантируем, что выходной путь — это наш патченный клон.
        tpl_doc.save(str(output_path))
        logger.info(
            "docx template path: cloned + patched {} paragraph(s) (coverage={:.2f})",
            patched, coverage,
        )
        return True
    except Exception as exc:
        logger.opt(exception=True).warning("docx template write failed: {}", exc)
        return False
