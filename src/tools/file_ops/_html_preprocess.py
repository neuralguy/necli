"""HTML-препроцессинг для create_docx через html.parser (надёжнее regex).

Функции:
  - wrap_table_cells: оборачивает inline-содержимое <td>/<th> в <p>,
    чтобы pandoc корректно парсил формулы и inline-форматирование внутри.

(apply_text_align удалён — выравнивание обрабатывается через
_extract_styled_spans + _apply_styles_to_docx в docx_writer post-process.)
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from io import StringIO

_BLOCK_TAGS = {"p", "div", "ul", "ol", "table", "blockquote", "pre",
               "h1", "h2", "h3", "h4", "h5", "h6"}
_VOID_TAGS = {"br", "hr", "img", "meta", "link", "input", "col"}


class _CellWrapper(HTMLParser):
    """Оборачивает содержимое td/th в <p> если внутри нет блочных тегов."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=False)
        self.out = StringIO()
        self._in_cell = False
        self._cell_buf = StringIO()
        self._cell_depth = 0
        self._cell_tag = ""
        self._cell_has_block = False

    def handle_starttag(self, tag, attrs):
        tag_l = tag.lower()
        if tag_l in ("td", "th") and not self._in_cell:
            self._in_cell = True
            self._cell_tag = tag_l
            self._cell_buf = StringIO()
            self._cell_depth = 1
            self._cell_has_block = False
            self.out.write(self._fmt_tag(tag, attrs, self_closing=False))
            return
        if self._in_cell:
            if tag_l in ("td", "th"):
                self._cell_depth += 1
            if tag_l in _BLOCK_TAGS:
                self._cell_has_block = True
            self._cell_buf.write(self._fmt_tag(tag, attrs, self_closing=False))
            return
        self.out.write(self._fmt_tag(tag, attrs, self_closing=False))

    def handle_startendtag(self, tag, attrs):
        if self._in_cell:
            self._cell_buf.write(self._fmt_tag(tag, attrs, self_closing=True))
        else:
            self.out.write(self._fmt_tag(tag, attrs, self_closing=True))

    def handle_endtag(self, tag):
        tag_l = tag.lower()
        if self._in_cell and tag_l == self._cell_tag:
            self._cell_depth -= 1
            if self._cell_depth == 0:
                content = self._cell_buf.getvalue()
                stripped = content.strip()
                if stripped and not self._cell_has_block:
                    self.out.write("<p>")
                    self.out.write(content)
                    self.out.write("</p>")
                else:
                    self.out.write(content)
                self.out.write(f"</{tag}>")
                self._in_cell = False
                self._cell_tag = ""
                return
        if self._in_cell:
            self._cell_buf.write(f"</{tag}>")
            return
        self.out.write(f"</{tag}>")

    def handle_data(self, data):
        if self._in_cell:
            self._cell_buf.write(data)
        else:
            self.out.write(data)

    def handle_entityref(self, name):
        chunk = f"&{name};"
        if self._in_cell:
            self._cell_buf.write(chunk)
        else:
            self.out.write(chunk)

    def handle_charref(self, name):
        chunk = f"&#{name};"
        if self._in_cell:
            self._cell_buf.write(chunk)
        else:
            self.out.write(chunk)

    def handle_comment(self, data):
        chunk = f"<!--{data}-->"
        if self._in_cell:
            self._cell_buf.write(chunk)
        else:
            self.out.write(chunk)

    @staticmethod
    def _fmt_tag(tag, attrs, *, self_closing: bool) -> str:
        parts = [f"<{tag}"]
        for k, v in attrs:
            if v is None:
                parts.append(f" {k}")
            else:
                v_esc = v.replace('"', '"')
                parts.append(f' {k}="{v_esc}"')
        if self_closing or tag.lower() in _VOID_TAGS:
            parts.append(" />")
        else:
            parts.append(">")
        return "".join(parts)


_KNOWN_HTML_TAGS = {
    "html", "head", "body", "meta", "link", "title", "style", "script",
    "p", "div", "span", "br", "hr", "a", "img",
    "h1", "h2", "h3", "h4", "h5", "h6",
    "ul", "ol", "li", "dl", "dt", "dd",
    "table", "thead", "tbody", "tfoot", "tr", "td", "th", "caption",
    "colgroup", "col",
    "strong", "b", "em", "i", "u", "s", "strike", "del", "ins", "mark",
    "sub", "sup", "small", "big", "code", "pre", "kbd", "samp", "var", "tt",
    "blockquote", "q", "cite", "abbr", "address", "bdo", "bdi",
    "figure", "figcaption", "section", "article", "header", "footer",
    "nav", "aside", "main", "details", "summary",
    # HTML5-теги, которые валидны, но раньше выпадали в stray-стэш и терялись.
    "wbr", "video", "audio", "source", "track", "picture",
    "time", "template", "data", "dialog", "menu", "fieldset",
    "legend", "label", "button", "select", "option", "optgroup",
    "textarea", "form", "datalist", "output", "progress", "meter",
    "ruby", "rt", "rp", "canvas", "svg", "math", "iframe", "embed",
    "object", "param", "map", "area", "noscript", "base",
}



_ANGLE_TOKEN_RE = re.compile(r"</?\s*([^<>\s/]+)[^<>]*>")
_ANGLE_ENTITY_RE = re.compile(r"</?\s*([^&<>\s/]+)[^&]*?>")
_STRAY_PH_PREFIX = "NECLISTRAYANGLE"
_STRAY_PH_RE = re.compile(rf"{_STRAY_PH_PREFIX}([0-9]+)X")


def extract_stray_angles(html: str) -> tuple[str, list[str]]:
    """Выносит литеральные/сущностные `<...>` (не HTML-теги) в стэш, ставит плейсхолдер.

    pandoc HTML reader трактует `<программа>`, `<оператор>`, `<выр>` (BNF-
    нетерминалы, плейсхолдеры в тексте) как неизвестные теги и МОЛЧА выкидывает
    их. Никакое экранирование не проходит надёжно через doctype-обёртку, поэтому,
    как и с кодом, не отдаём их в pandoc вообще: заменяем на буквенный
    плейсхолдер, а после конвертации впечатываем исходный текст обратно в runs
    через restore_stray_angles.

    Ловим ДВЕ формы:
    1. Литеральную `<имя…>`.
    2. Сущностную `<имя…>` — wrap_table_cells (HTMLParser) перекодирует
       литеральные угловые скобки в сущности, поэтому к моменту вызова в реальном
       pipeline BNF-нетерминалы выглядят именно так. В стэш кладём ДЕкодированный
       текст (`<имя>`), чтобы restore впечатал нормальные скобки.

    Возвращает (html_с_плейсхолдерами, [raw_text, ...]).

    Вызывать ПОСЛЕ extract_code_blocks (код в стэше) и ПОСЛЕ wrap_table_cells.
    """
    import html as _htmlmod

    stash: list[str] = []

    def _repl_literal(m: re.Match[str]) -> str:
        name = m.group(1).lower()
        if name in _KNOWN_HTML_TAGS:
            return m.group(0)
        idx = len(stash)
        stash.append(m.group(0))
        return f"{_STRAY_PH_PREFIX}{idx}X"

    def _repl_entity(m: re.Match[str]) -> str:
        name = m.group(1).lower()
        if name in _KNOWN_HTML_TAGS:
            return m.group(0)
        idx = len(stash)
        stash.append(_htmlmod.unescape(m.group(0)))
        return f"{_STRAY_PH_PREFIX}{idx}X"

    out = _ANGLE_TOKEN_RE.sub(_repl_literal, html)
    out = _ANGLE_ENTITY_RE.sub(_repl_entity, out)
    return out, stash




def restore_stray_angles(docx_path, stash: list[str]) -> None:
    """Впечатывает литеральные `<...>` вместо плейсхолдеров в готовом .docx.

    Идёт по всем runs (включая таблицы), заменяет вхождения плейсхолдера на
    исходный текст. Может быть несколько плейсхолдеров в одном run.
    """
    if not stash:
        return

    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    def _iter_paragraphs():
        yield from doc.paragraphs
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    changed = False
    for para in _iter_paragraphs():
        if _STRAY_PH_PREFIX not in para.text:
            continue
        for run in para.runs:
            txt = run.text or ""
            if _STRAY_PH_PREFIX not in txt:
                continue

            def _sub(m: re.Match[str]) -> str:
                idx = int(m.group(1))
                return stash[idx] if 0 <= idx < len(stash) else m.group(0)

            new_txt = _STRAY_PH_RE.sub(_sub, txt)
            if new_txt != txt:
                t_el = run._element.find(qn("w:t"))
                if t_el is not None:
                    t_el.text = new_txt
                    t_el.set(qn("xml:space"), "preserve")
                else:
                    run.text = new_txt
                changed = True

    if changed:
        doc.save(str(docx_path))


_PRE_CODE_MASK_RE = re.compile(r"<pre\b[^>]*>.*?</pre>", flags=re.IGNORECASE | re.DOTALL)


def wrap_table_cells(html: str) -> str:
    # HTMLParser лучше не пускать внутрь <pre>: он трактует `<Lexeme>` как
    # настоящий тег (приводит к нижнему регистру, перекраивает атрибуты) и
    # декодирует уже выставленные нами сущности обратно. Маскируем code-блоки
    # на время парсинга и возвращаем дословно.
    stash: list[str] = []

    def _mask(m: re.Match[str]) -> str:
        stash.append(m.group(0))
        return f"\x00PRE{len(stash) - 1}\x00"

    masked = _PRE_CODE_MASK_RE.sub(_mask, html)
    p = _CellWrapper()
    p.feed(masked)
    out = p.out.getvalue()
    for i, block in enumerate(stash):
        out = out.replace(f"\x00PRE{i}\x00", block)
    return out




_CODE_BLOCK_RE = re.compile(
    r"<pre\b[^>]*>\s*<code\b[^>]*>(.*?)</code>\s*</pre>",
    flags=re.IGNORECASE | re.DOTALL,
)
_INLINE_CODE_RE = re.compile(
    r"<code\b[^>]*>(.*?)</code>",
    flags=re.IGNORECASE | re.DOTALL,
)
# Уникальный плейсхолдер-токен. Состоит из букв/цифр — pandoc проносит его
# через HTML reader дословно (не трогает как тег и не разбивает на слова).
_CODE_PH_PREFIX = "NECLICODEBLOCK"
_CODE_PH_RE = re.compile(rf"{_CODE_PH_PREFIX}([0-9]+)X")


def _decode_entities(text: str) -> str:
    """Декодирует HTML-сущности в исходный текст кода (для хранения «как есть»)."""
    import html as _html
    return _html.unescape(text)


def extract_code_blocks(html: str) -> tuple[str, list[tuple[str, str]]]:
    """Выносит тело каждого <pre><code> и inline <code> в стэш, ставит плейсхолдер.

    Возвращает (html_с_плейсхолдерами, [(kind, raw_text), ...]), где
    kind ∈ {"block", "inline"}.

    Зачем: pandoc HTML reader ВСЕГДА деодирует сущности внутри <code> и
    выкидывает tag-подобные токены (`<iostream>`, `vector<Lexeme>`,
    `<программа>`) — даже двойное экранирование не спасает. Поэтому код вообще
    не отдаём в pandoc: вместо тела ставим буквенный плейсхолдер, а после
    конвертации впечатываем сырой текст напрямую в runs через python-docx
    (restore_code_blocks).
    """
    stash: list[tuple[str, str]] = []

    def _block(m: re.Match[str]) -> str:
        idx = len(stash)
        stash.append(("block", _decode_entities(m.group(1))))
        return f"<pre><code>{_CODE_PH_PREFIX}{idx}X</code></pre>"

    def _inline(m: re.Match[str]) -> str:
        body = m.group(1)
        # Не трогаем <code>, чьё тело — уже выставленный нами плейсхолдер
        # (он остался от блочной замены <pre><code>…</code></pre>).
        if _CODE_PH_RE.fullmatch(body.strip()):
            return m.group(0)
        idx = len(stash)
        stash.append(("inline", _decode_entities(body)))
        return f"<code>{_CODE_PH_PREFIX}{idx}X</code>"

    html = _CODE_BLOCK_RE.sub(_block, html)
    html = _INLINE_CODE_RE.sub(_inline, html)
    return html, stash


def restore_code_blocks(docx_path, stash: list[tuple[str, str]]) -> None:
    """Впечатывает сырой код вместо плейсхолдеров в готовом .docx.

    Идёт по всем параграфам (в т.ч. в таблицах). Где текст параграфа содержит
    плейсхолдер — заменяет соответствующий run на сырой текст, многострочный
    код разбивается на строки через <w:br/>. Шрифт/стиль run'а (Source Code /
    inline `code`) сохраняется.
    """
    if not stash:
        return

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    def _iter_paragraphs():
        yield from doc.paragraphs
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def _force_monospace(run) -> None:
        """Принудительно ставит Courier New на run.

        pandoc вешает на код символьный стиль VerbatimChar без шрифта; в
        reference.docx его нет, поэтому run наследует TNR из docDefaults и
        код рендерится пропорциональным шрифтом. Прописываем rFonts прямо в
        run, чтобы шрифт не зависел от наличия/настройки символьного стиля.
        """
        r_el = run._element
        rpr = r_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            r_el.insert(0, rpr)
        for old in rpr.findall(qn("w:rFonts")):
            rpr.remove(old)
        rfonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), "Courier New")
        rpr.append(rfonts)

    def _set_run_multiline(run, text: str) -> None:
        r_el = run._element
        for child in list(r_el):
            if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab"), qn("w:cr")):
                r_el.remove(child)
        _force_monospace(run)
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                r_el.append(OxmlElement("w:br"))
            if line:
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = line
                r_el.append(t_el)

    for para in _iter_paragraphs():
        if _CODE_PH_PREFIX not in para.text:
            continue
        for run in para.runs:
            m = _CODE_PH_RE.search(run.text or "")
            if not m:
                continue
            idx = int(m.group(1))
            if 0 <= idx < len(stash):
                _set_run_multiline(run, stash[idx][1])

    doc.save(str(docx_path))
